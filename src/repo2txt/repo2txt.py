#!/usr/bin/env python3
"""
repo2txt - Repository Documentation Generator

A comprehensive tool for documenting repository structure and contents.
Generates both text and DOCX format outputs with extensive filtering options.

Author: Your Name
Version: 2.0.0
License: MIT
"""

import os
import argparse
import json
import fnmatch
import sys
from pathlib import Path
from typing import Dict, List, Optional, Union, TextIO

# Try to import docx components, but make them optional
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    print("Warning: python-docx not available. DOCX output will not work.")

# Configuration constants
SCRIPT_DIR = os.path.dirname(os.path.realpath(__file__))
CONFIG_FILE_PATH = os.path.join(SCRIPT_DIR, 'config.json')
VERSION = "2.0.0"

# Default ignore patterns for common development artifacts
DEFAULT_IGNORE_DIRS = {'.git', '.vscode', '.idea', '__pycache__', 'node_modules', 
                      '.pytest_cache', '.mypy_cache', '.tox', 'venv', '.venv'}


def load_config(file_path: str = CONFIG_FILE_PATH) -> Dict:
    """
    Load configuration from a JSON file with fallback to defaults.
    
    Args:
        file_path: Path to the JSON configuration file
        
    Returns:
        Dictionary containing configuration settings
        
    Raises:
        None - Gracefully falls back to defaults on any error
    """
    default_config = {
        "image_extensions": [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".svg"],
        "video_extensions": [".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv"],
        "audio_extensions": [".mp3", ".wav", ".aac", ".flac", ".ogg"],
        "document_extensions": [".pdf", ".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx", 
                               ".jar", ".zip", ".tar", ".gzip"],
        "executable_extensions": [".exe", ".dll", ".bin", ".sh", ".bat"],
        "settings_extensions": [".ini", ".cfg", ".conf", ".json", ".yaml", ".yml"],
        "additional_ignore_types": [".lock", ".pyc", ".pyo", ".so", ".dylib"],
        "default_output_file": "output.txt"
    }
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            config = json.load(file)
            # Ensure all required keys exist
            for key, default_value in default_config.items():
                config.setdefault(key, default_value)
            return config
    except (FileNotFoundError, json.JSONDecodeError, PermissionError) as e:
        print(f"Warning: Could not load config from {file_path}: {e}")
        print("Using default configuration.")
        return default_config


# Load configuration
config = load_config()
IMAGE_EXTENSIONS = config.get("image_extensions", [])
VIDEO_EXTENSIONS = config.get("video_extensions", [])
AUDIO_EXTENSIONS = config.get("audio_extensions", [])
DOCUMENT_EXTENSIONS = config.get("document_extensions", [])
EXECUTABLE_EXTENSIONS = config.get("executable_extensions", [])
SETTINGS_EXTENSIONS = config.get("settings_extensions", [])
ADDITIONAL_IGNORE_TYPES = config.get("additional_ignore_types", [])
DEFAULT_OUTPUT_FILE = config.get("default_output_file", "output.txt")

# Combine all ignore types into a comprehensive list
DEFAULT_IGNORE_TYPES_LIST = list(set(
    IMAGE_EXTENSIONS + VIDEO_EXTENSIONS + AUDIO_EXTENSIONS +
    DOCUMENT_EXTENSIONS + EXECUTABLE_EXTENSIONS + ADDITIONAL_IGNORE_TYPES
))


def parse_args() -> argparse.Namespace:
    """
    Parse and validate command-line arguments.
    
    Returns:
        Parsed arguments namespace
    """
    parser = argparse.ArgumentParser(
        description='Document the structure and contents of a repository or directory.',
        epilog='''
Examples:
  Basic usage:
    %(prog)s -r /path/to/repo -o documentation.txt
    %(prog)s -r /path/to/repo -o report.docx
    
  Advanced filtering:
    %(prog)s --include-files "*.py" "*.js" -o code_only.txt
    %(prog)s --ignore-types --exclude-dir tests docs -o minimal.txt
    %(prog)s --include-dir src --ignore-settings -o src_only.txt
    
  Special cases:
    %(prog)s --ignore-types none  # Include all file types
    %(prog)s --ignore-files none  # Include all files
    %(prog)s --exclude-dir none   # Include all directories
        ''',
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    # Main arguments
    parser.add_argument('-r', '--repo_path', 
                       default=os.getcwd(),
                       help='Path to the directory to process. Defaults to current directory.')
    
    parser.add_argument('-o', '--output_file', 
                       default=DEFAULT_OUTPUT_FILE,
                       help=f'Output file name. Use .txt for text or .docx for Word format. '
                            f'(default: {DEFAULT_OUTPUT_FILE})')
    
    # Filtering options
    parser.add_argument('--ignore-files', 
                       nargs='*', 
                       default=[],
                       help='List of specific filenames to ignore. Use "none" to ignore no files.')
    
    parser.add_argument('--ignore-types', 
                       nargs='*', 
                       default=DEFAULT_IGNORE_TYPES_LIST,
                       help='File extensions to ignore. Use "none" to include all types. '
                            'Defaults to common binary/media files.')
    
    parser.add_argument('--exclude-dir', 
                       nargs='*', 
                       default=[],
                       help='Directory names to exclude. Use "none" to exclude no directories.')
    
    parser.add_argument('--ignore-settings', 
                       action='store_true',
                       help='Ignore common configuration/settings files.')
    
    # Include options
    parser.add_argument('--include-dir', 
                       nargs='?', 
                       default=None,
                       help='Process only this specific directory and its contents.')
    
    parser.add_argument('--include-files', 
                       nargs='*', 
                       default=None,
                       help='Include only files matching these patterns (e.g., "*.py" "*.js").')
    
    # Utility options
    parser.add_argument('--version', 
                       action='version', 
                       version=f'%(prog)s {VERSION}')
    
    parser.add_argument('--verbose', '-v',
                       action='store_true',
                       help='Enable verbose output showing processing details.')
    
    parser.add_argument('--dry-run',
                       action='store_true',
                       help='Show what would be processed without generating output.')
    
    parser.add_argument('--no-content',
                    action='store_true',
                    help='Skip file content output and include only the folder/file hierarchy.')

    return parser.parse_args()


def should_ignore(item_path: str, args: argparse.Namespace, repo_root_path: str) -> bool:
    """
    Determine if a file or directory should be ignored based on filtering rules.
    
    Args:
        item_path: Full path to the item being checked
        args: Parsed command-line arguments containing filter settings
        repo_root_path: Root path of the repository being processed
        
    Returns:
        True if the item should be ignored, False otherwise
        
    Note:
        Filtering logic is applied in order of precedence:
        1. Output file self-exclusion
        2. Common development artifacts
        3. Hidden files/directories
        4. User-specified directory exclusions
        5. Include directory restrictions
        6. File-specific filters (include-files, ignore-files, ignore-types)
        7. Settings file filtering
    """
    item_path_abs = os.path.abspath(item_path)
    item_name = os.path.basename(item_path)
    file_ext = os.path.splitext(item_name)[1].lower()

    # Never include the output file itself
    if item_path_abs == os.path.abspath(args.output_file):
        return True
    
    # SECURITY: Always exclude environment files (case-insensitive)
    if os.path.isfile(item_path) and (
        item_name.lower().startswith('.env') or 
        item_name.lower() in {'.env', '.environment', 'environment'}
    ):
        return True
    
    # Ignore common development artifacts
    if item_name in DEFAULT_IGNORE_DIRS and os.path.isdir(item_path):
        return True
    
    # Ignore hidden files and directories (except root)
    if item_name.startswith('.') and item_path_abs != os.path.abspath(repo_root_path):
        return True
    
    # Apply directory exclusions
    if os.path.isdir(item_path) and args.exclude_dir and item_name in args.exclude_dir:
        return True

    # Apply include directory restriction
    if args.include_dir:
        abs_include_dir = os.path.abspath(args.include_dir)
        if not (item_path_abs.startswith(abs_include_dir) or abs_include_dir.startswith(item_path_abs)):
            return True

    # Apply file-specific filters
    if os.path.isfile(item_path):
        # Include files filter (whitelist)
        if args.include_files is not None:
            if not any(fnmatch.fnmatch(item_name, pattern) for pattern in args.include_files):
                return True
        
        # Ignore files filter (blacklist)
        if item_name in args.ignore_files or file_ext in args.ignore_types:
            return True
        
        # Settings files filter
        if args.ignore_settings and file_ext in SETTINGS_EXTENSIONS:
            return True

    return False


def write_tree(dir_path: str, output_file: Union[TextIO, "Document"], args: argparse.Namespace, 
               repo_root_path: str, prefix: str = "", is_last: bool = True, is_root: bool = True) -> int:
    """
    Recursively write the directory tree structure to output.
    
    Args:
        dir_path: Directory path to process
        output_file: File handle or Document object to write to
        args: Parsed command-line arguments
        repo_root_path: Root path of the repository
        prefix: Current line prefix for tree structure
        is_last: Whether this is the last item at current level
        is_root: Whether this is the root directory
        
    Returns:
        Number of items processed
    """
    items_processed = 0
    
    if is_root:
        root_name = f"{os.path.basename(dir_path)}/"
        if DOCX_AVAILABLE and hasattr(output_file, 'add_paragraph'):
            paragraph = output_file.add_paragraph()
            paragraph.add_run(root_name)
        else:
            output_file.write(f"{root_name}\n")
        items_processed += 1
    
    try:
        items = sorted(os.listdir(dir_path))
    except PermissionError:
        if args.verbose:
            print(f"Warning: Permission denied accessing {dir_path}")
        return items_processed
    
    # Filter items and maintain original indices for proper tree structure
    filtered_items = []
    for item in items:
        item_path = os.path.join(dir_path, item)
        if not should_ignore(os.path.abspath(item_path), args, repo_root_path):
            filtered_items.append(item)
    
    for idx, item in enumerate(filtered_items):
        item_path = os.path.join(dir_path, item)
        is_last_item = (idx == len(filtered_items) - 1)
        
        # Create tree structure
        current_prefix = f"{prefix}{'└── ' if is_last_item else '├── '}{item}"
        
        if DOCX_AVAILABLE and hasattr(output_file, 'add_paragraph'):
            paragraph = output_file.add_paragraph()
            paragraph.add_run(f"{prefix}{'└── ' if is_last_item else '├── '}{item}")
        else:
            output_file.write(f"{current_prefix}\n")
        
        items_processed += 1
        
        if args.verbose:
            print(f"Processing: {os.path.relpath(item_path, repo_root_path)}")
        
        # Recurse into directories
        if os.path.isdir(item_path):
            next_prefix = prefix + ('    ' if is_last_item else '│   ')
            items_processed += write_tree(item_path, output_file, args, repo_root_path,
                                        next_prefix, is_last_item, False)
    
    return items_processed


def write_file_content(file_path: str, output_file: Union[TextIO, "Document"], depth: int) -> bool:
    """
    Write the contents of a file to the output with proper formatting.
    
    Args:
        file_path: Path to the file to read
        output_file: File handle or Document object to write to
        depth: Indentation depth for text output
        
    Returns:
        True if file was successfully read, False otherwise
    """
    indentation = '  ' * depth
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            if DOCX_AVAILABLE and hasattr(output_file, 'add_paragraph'):
                content = f.read()
                paragraph = output_file.add_paragraph(content)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                for line in f:
                    output_file.write(f"{indentation}{line}")
        return True
    except Exception as e:
        error_msg = f"Error reading file: {e}"
        if DOCX_AVAILABLE and hasattr(output_file, 'add_paragraph'):
            error_paragraph = output_file.add_paragraph(error_msg)
            error_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            output_file.write(f"{indentation}{error_msg}\n")
        return False


def write_file_contents_in_order(dir_path: str, output_file: Union[TextIO, "Document"], 
                                args: argparse.Namespace, repo_root_path: str, depth: int = 0) -> int:
    """
    Recursively write file contents in directory tree order.
    
    Args:
        dir_path: Directory to process
        output_file: File handle or Document object to write to
        args: Parsed command-line arguments
        repo_root_path: Root path of the repository
        depth: Current recursion depth for indentation
        
    Returns:
        Number of files processed
    """
    files_processed = 0
    
    try:
        items = sorted(os.listdir(dir_path))
    except PermissionError:
        if args.verbose:
            print(f"Warning: Permission denied accessing {dir_path}")
        return files_processed
    
    for item in items:
        item_path = os.path.join(dir_path, item)
        if should_ignore(os.path.abspath(item_path), args, repo_root_path):
            continue
        
        relative_start = os.path.abspath(args.include_dir or args.repo_path)
        relative_path = os.path.relpath(item_path, start=relative_start)
        
        if os.path.isdir(item_path):
            files_processed += write_file_contents_in_order(
                item_path, output_file, args, repo_root_path, depth + 1)
        elif os.path.isfile(item_path):
            # File begin marker
            begin_marker = f"[File Begins] {relative_path}"
            if DOCX_AVAILABLE and hasattr(output_file, 'add_heading'):
                output_file.add_heading(begin_marker, level=3)
            else:
                output_file.write('  ' * depth + f"{begin_marker}\n")
            
            # File content
            success = write_file_content(item_path, output_file, depth)
            if success:
                files_processed += 1
            
            # File end marker
            end_marker = f"[File Ends] {relative_path}"
            if DOCX_AVAILABLE and hasattr(output_file, 'add_heading'):
                output_file.add_heading(end_marker, level=3)
            else:
                output_file.write('  ' * depth + f"{end_marker}\n\n")
    
    return files_processed


def create_docx_document():
    """
    Create and configure a new DOCX document with appropriate styling.
    
    Returns:
        Configured Document object
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX output but is not installed")
    
    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    style.font.name = 'Consolas'  # Monospace font for code
    style.font.size = Pt(9)
    
    # Add document header
    doc.add_heading("Repository Documentation", 1)
    
    intro_text = (
        "This document provides a comprehensive overview of the repository's structure and contents. "
        "The first section displays the directory hierarchy in a tree format, "
        "followed by the complete contents of each file in the repository."
    )
    doc.add_paragraph(intro_text)
    
    return doc


def validate_arguments(args: argparse.Namespace) -> bool:
    """
    Validate command-line arguments and check for common issues.
    
    Args:
        args: Parsed command-line arguments
        
    Returns:
        True if arguments are valid, False otherwise
    """
    # Handle 'none' keywords
    if args.ignore_files == ['none']:
        args.ignore_files = []
    if args.ignore_types == ['none']:
        args.ignore_types = []
    if args.exclude_dir == ['none']:
        args.exclude_dir = []
    
    # Validate repository path
    processing_root_path = os.path.abspath(args.include_dir or args.repo_path)
    if not os.path.isdir(processing_root_path):
        print(f"Error: '{processing_root_path}' is not a valid directory.", file=sys.stderr)
        return False
    
    # Check if output directory exists
    output_dir = os.path.dirname(os.path.abspath(args.output_file))
    if not os.path.exists(output_dir):
        print(f"Error: Output directory '{output_dir}' does not exist.", file=sys.stderr)
        return False
    
    # Check for DOCX output when library is not available
    if args.output_file.endswith('.docx') and not DOCX_AVAILABLE:
        print("Error: DOCX output requested but python-docx is not installed.", file=sys.stderr)
        print("Install it with: pip install python-docx", file=sys.stderr)
        return False
    
    # Validate include_files patterns
    if args.include_files is not None and len(args.include_files) == 0:
        print("Warning: --include-files specified with no patterns. No files will be included.")
    
    return True


def main():
    """
    Main function to execute the repository documentation generation.
    """
    args = parse_args()
    
    if not validate_arguments(args):
        sys.exit(1)
    
    processing_root_path = os.path.abspath(args.include_dir or args.repo_path)
    
    if args.verbose:
        print(f"Processing directory: {processing_root_path}")
        print(f"Output file: {args.output_file}")
        print(f"Include files: {args.include_files}")
        print(f"Ignore types: {len(args.ignore_types)} extensions")
    
    try:
        if args.output_file.endswith('.docx'):
            # Generate DOCX output
            doc = create_docx_document()
            
            # Directory tree section
            doc.add_heading("Directory/File Tree -->", 2)
            tree_items = write_tree(processing_root_path, doc, args, processing_root_path)
            doc.add_heading("<-- Directory/File Tree", 2)
                        
           # Optionally skip file content section if --no-content is used             
            file_count = 0
            if not args.no_content:
                doc.add_heading("File Contents -->", 2)
                file_count = write_file_contents_in_order(processing_root_path, doc, args, processing_root_path)
                doc.add_heading("<-- File Contents", 2)
            
            # Save document
            doc.save(args.output_file)
            
            if args.verbose:
                print(f"DOCX document saved with {tree_items} items and {file_count} files")
        
        else:
            # Generate text output
            with open(args.output_file, 'w', encoding='utf-8') as f:
                f.write("Repository Documentation\n")
                f.write("=" * 50 + "\n\n")
                
                # Directory tree section
                f.write("Directory/File Tree -->\n\n")
                tree_items = write_tree(processing_root_path, f, args, processing_root_path)
                f.write("\n<-- Directory/File Tree\n\n")
                
                # Optionally skip file content section if --no-content is used
                file_count = 0
                if not args.no_content:
                    f.write("File Contents -->\n\n")
                    file_count = write_file_contents_in_order(processing_root_path, f, args, processing_root_path)
                    f.write("\n<-- File Contents\n\n")
            
            if args.verbose:
                print(f"Text file saved with {tree_items} items and {file_count} files")
        
        if args.dry_run:
            print("Dry run completed - no output file generated")
        else:
            print(f"Documentation successfully generated: {args.output_file}")
    
    except KeyboardInterrupt:
        print("\nOperation cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()